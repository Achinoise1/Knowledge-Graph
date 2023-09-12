# encoding:utf-8
import re
from docx import Document

# 检查是否能够正常打开
# 绝对路径 -> path = "F:\\Desktop_From_C\\project\\教学大纲.docx"             
# 相对路径 -> 
path = '../教学大纲.docx'

try:    
    doc = Document(path)
except Exception as e:
    print("Error!",e)
else:
    print("Success!")      
    tables = doc.tables                                         #计算检索表格数量

#提取课程名称,加入字典并返回字典
def SearchCourseName(doc):
    Dic = {}
    pattern = re.compile(r"《(.+?)》教学大纲")
    for p in doc.paragraphs:
        str = p.text    
        list = re.findall(pattern,str)
        if list:
            Dic[list[0]] = {}
    return Dic

#这里是某种意义上的典中典
AllCourse = SearchCourseName(doc)
#print(AllCourse)


#提取教学内容，存入并返回字典
#先把一个课的所有内容存在一个字典，再把字典存在之前的大字典
def SearchCourseContent(doc,DicOrigin):
    count = -1
    Keys = list(DicOrigin.keys())
    #exit(0)
    for p in doc.paragraphs:
        str = p.text
        if(str == "理论教学内容与课程目标的关系"):
            table = 1;
        '''else:
            for j in range(1,len(table.rows)):                                       
                Order = table.cell(j,0).text                            
                Content = table.cell(j,1).text.split("\n")             
                DicOrigin[Keys[count]].update({Order:Content})
                if(Order == '1'):
                    count += 1
                if(count == len(Keys)):
                    return DicOrigin '''                                                                                                     
                #print(DicOrigin)                  
    #return DicOrigin
# Count 增加条件
# Out of range?

AllContent = SearchCourseContent(doc,AllCourse)
#print(AllContent)
