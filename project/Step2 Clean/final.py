import io
import sys
from base64 import encode
import encodings
import docx
import re
import json
from docx import Document
from torch import count_nonzero

sys.stdout = io.TextIOWrapper(sys.stdout.buffer,encoding='utf-8')
# 检查是否能够正常打开
path = "F:\\Desktop_From_C\\project\\教学大纲.docx"             
try:    
    doc = Document(path)
except Exception as e:
    print("Error!",e)
else:
    print("Success!")      
    tables = doc.tables                                         #计算检索表格数量
    TableLen = len(tables)

#提取课程名称,加入字典并返回字典
def SearchCourseName(doc):
    Dic = {}
    #pattern = re.compile(r"《(.+?[^*])》教学大纲")         #用[^*]筛掉英文的大纲
    pattern = re.compile(r"《(.+?)》教学大纲") 
    for p in doc.paragraphs:
        str = p.text    
        list = re.findall(pattern,str)
        if list:
            Dic[list[0]] = {}
    return Dic

#这里是某种意义上的典中典
AllCourse = SearchCourseName(doc)
#print(AllCourse)

#提取教学内容，存入并返回字典
#先把一个课的所有内容存在一个字典，再把字典存在之前的大字典
def SearchCourseContent(doc,DicOrigin):
    count = -1
    Keys = list(DicOrigin.keys())
    for i in range(1,len(doc.tables)):                                           
        table = tables[i]             
        if((table.cell(0,0).text != "序号" and table.cell(0,0).text != "序\n号" and table.cell(0,0).text != "Num.") or ((table.cell(1,1).text[0] == "实" and table.cell(1,1).text[1] == "验") or (table.cell(1,1).text[0] == 'L' and table.cell(1,1).text[1] == 'a' and table.cell(1,1).text[2] == 'b')) ):
            continue
        else:
            if(count!=-1 and Keys[count]=="软件体系结构"):
                count += 1
            for j in range(1,len(table.rows)):                                       
                Order = table.cell(j,0).text                        
                Content = re.split('[\n]',table.cell(j,1).text.strip())                    #Content = table.cell(j,1).text.split("\n")                            
                if(Order == '1'):
                    count += 1
                if(count == len(Keys)):
                    return DicOrigin                
                DicOrigin[Keys[count]].update({Order:Content})
    DicOrigin.update(
                    {
                        "大数据应用项目实践": {
                            "1": ['Linux基本操作','了解Linux基本概念及应用','掌握基于虚拟机的Linux安装配置','掌握虚拟机网络配置','掌握Linux基本操作及Shell脚本编写'],
                            "2": ['Hadoop生态及其安装','讨论Hadoop生态系统','Hadoop集群搭建'],
                            "3": ['HDFS分布式文件系统','HDFS基本概念','HDFS的组成部分及体系结构','HDFS的常用API操作','HDFS的工作机制。'],
                            "4": ['MapReduce分布式文件系统','MapReduce基本概念及Hadoop序列化','MapReduce框架搭建','用MapReduce实现数据集连接、数据查重、词频统计等操作。'],
                            "5": ['MapReduce企业应用实践 ','基于MapReduce的文件单词统计','基于MapReduce的手机流量信息挖掘分析'],
                            "6": ['综合应用实践 ','基于特定行业应用问题，开展需求分析、方案设计及编码实现','以小组为单位，完成项目实践并汇报']
                        }
                    }
                )                                                                                                 
    return DicOrigin


AllContent = SearchCourseContent(doc,AllCourse)
 
def Clean_zh(BigDic):
    KeyList = list(BigDic.keys())
    formatChap = re.compile(r"第([1-9][0-9]?|[一二三四五六七八九十]+)章") #or re.compile(r"Chapter[0-9]")
    formatTea = re.compile(r"教学内容[0-9]?")       #为啥没有问号匹配不到
    formatKnow = re.compile(r"知识单元[0-9]?[：]?")
    formatNum = re.compile(r"([0-9][0-9]?|[一二三四五六七八九十]+|[a-z])\.[0-9]?")
    formatSign = re.compile(r"[，；。：]$") #or re.compile(r"\s")
    formatSpace = re.compile(r"\s")   #formatSpace = re.compile(r"[^a-zA-Z]+\s")
    formatOrder = re.compile(r"(（[0-9][0-9]?）)")
    formatOrderE = re.compile(r"\([0-9][0-9]?\)")
    for i in range(0,len(KeyList)):
        if(KeyList[i][len(KeyList[i])-1]=='*'):
            continue
        OrderList = list(BigDic[KeyList[i]].keys())
        for j in range(0,len(OrderList)):
            ContentList = BigDic[KeyList[i]][OrderList[j]]
            for l in range(0,len(ContentList)):
                BigDic[KeyList[i]][OrderList[j]][l] = re.sub(formatChap,"",BigDic[KeyList[i]][OrderList[j]][l])
                BigDic[KeyList[i]][OrderList[j]][l] = re.sub(formatTea,"",BigDic[KeyList[i]][OrderList[j]][l])
                BigDic[KeyList[i]][OrderList[j]][l] = re.sub(formatKnow,"",BigDic[KeyList[i]][OrderList[j]][l])
                BigDic[KeyList[i]][OrderList[j]][l] = re.sub(formatNum,"",BigDic[KeyList[i]][OrderList[j]][l])
                BigDic[KeyList[i]][OrderList[j]][l] = re.sub(formatSign,"",BigDic[KeyList[i]][OrderList[j]][l])
                
                BigDic[KeyList[i]][OrderList[j]][l] = re.sub(formatOrder,"",BigDic[KeyList[i]][OrderList[j]][l])
                BigDic[KeyList[i]][OrderList[j]][l] = re.sub(formatOrderE,"",BigDic[KeyList[i]][OrderList[j]][l])
                BigDic[KeyList[i]][OrderList[j]][l] = re.sub(formatSpace,"",BigDic[KeyList[i]][OrderList[j]][l])
                
    return BigDic
 
def Clean_eg(BigDic):
    
    formatChap = re.compile(r"Chapter[\s][0-9]")
    formatOrder = re.compile(r"\([0-9][0-9]?\)") 
    formatSign = re.compile(r"[:,]$")  
    OrderList = list(BigDic["可视化编程技术B*"].keys())
    for j in range(0,len(OrderList)):
        ContentList = BigDic["可视化编程技术B*"][OrderList[j]]
        for l in range(0,len(ContentList)):
            BigDic["可视化编程技术B*"][OrderList[j]][l] = re.sub(formatChap,"",BigDic["可视化编程技术B*"][OrderList[j]][l])
            BigDic["可视化编程技术B*"][OrderList[j]][l] = re.sub(formatOrder,"",BigDic["可视化编程技术B*"][OrderList[j]][l])
            BigDic["可视化编程技术B*"][OrderList[j]][l] = BigDic["可视化编程技术B*"][OrderList[j]][l].strip(" ")
            
    OrderList = list(BigDic["人机交互技术A*"].keys())
    for j in range(0,len(OrderList)):
        ContentList = BigDic["人机交互技术A*"][OrderList[j]]
        for l in range(0,len(ContentList)):
            BigDic["人机交互技术A*"][OrderList[j]][l] = re.sub(formatChap,"",BigDic["人机交互技术A*"][OrderList[j]][l])
            BigDic["人机交互技术A*"][OrderList[j]][l] = re.sub(formatOrder,"",BigDic["人机交互技术A*"][OrderList[j]][l])
            BigDic["人机交互技术A*"][OrderList[j]][l] = re.sub(formatSign,"",BigDic["人机交互技术A*"][OrderList[j]][l])
            BigDic["人机交互技术A*"][OrderList[j]][l] = BigDic["人机交互技术A*"][OrderList[j]][l].strip(" ")
    #    
    return BigDic 
           
AllContent = Clean_zh(AllContent)
AllCentent = Clean_eg(AllContent)


with open("a.json", "w",encoding='utf-8') as f:
    f.write(json.dumps(AllContent, ensure_ascii=False, indent=4, separators=(',', ':')))
