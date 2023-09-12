import docx
import re
from docx import Document

# 检查是否能够正常打开
path = "F:\\Desktop_From_C\\project\\教学大纲.docx"
# 相对路径 -> path = "..\教学大纲.docx"

''' 
副本？
查找文件中的点出现的位置
idx = path.rfind(".")
print(file1_name[0:idx])
print(file1_name[idx:])
利用切片添加副本文字
file = path[0:idx]+ " - 副本" + path[idx:]
'''

try:    
    doc = Document(path)
except Exception as e:
    print("Error!",e)
else:
    print("Success!")
    #计算检索表格数量   
    tables = doc.tables
    TableLen = len(tables)

#提取课程名称,加入字典并返回字典
def SearchCourseName(doc):
    Dic = {}
    pattern = re.compile(r"《(.+?)》教学大纲")
    for p in doc.paragraphs:
        str = p.text    
        list = re.findall(pattern,str)
        if list:
            Dic[list[0]] = {}
            #print(list)
    return Dic

#这里是列表
AllCourse = SearchCourseName(doc)
#print(AllCourse)


#提取教学内容，存入并返回字典
#先把一个课的所有内容存在一个字典，再把字典存在之前的大字典
def SearchCourseContent(Dic1):
    DicTemp = {}
    for i in range(1,20):                           #srds,第0个表格是哪个
        table = tables[i]
        for j in range(1,len(table.rows)):
            if(table.cell(0,0).text!="序号" and table.cell(0,0).text!="序\n号"):
                break;
            Order = table.cell(j,0).text                            
            Content = table.cell(j,1).text.split("\n")
            DicTemp.update({Order:Content})        
        print(DicTemp)
        
    return Dic1

AllContent = SearchCourseContent(AllCourse)
#print(AllContent)
