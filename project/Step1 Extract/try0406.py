import io
import sys
from base64 import encode
import encodings
import docx
import re
import json
from docx import Document

sys.stdout = io.TextIOWrapper(sys.stdout.buffer,encoding='utf-8')
# 检查是否能够正常打开
path = "F:\\Desktop_From_C\\project\\教学大纲.docx"             # 相对路径 -> path = "..\教学大纲.docx"
try:    
    doc = Document(path)
except Exception as e:
    print("Error!",e)
else:
    print("Success!")      
    tables = doc.tables                                         #计算检索表格数量
    TableLen = len(tables)



#提取课程名称,加入字典并返回字典
def SearchCourseName(doc):
    Dic = {}
    pattern = re.compile(r"《(.+?[^*])》教学大纲")         #用[^*]筛掉英文的大纲
    for p in doc.paragraphs:
        str = p.text    
        list = re.findall(pattern,str)
        if list:
            Dic[list[0]] = {}
    return Dic

#这里是某种意义上的典中典
AllCourse = SearchCourseName(doc)
#print(AllCourse)


#提取教学内容，存入并返回字典
#先把一个课的所有内容存在一个字典，再把字典存在之前的大字典
def SearchCourseContent(doc,DicOrigin):
    count = -1
    Keys = list(DicOrigin.keys())
    for i in range(1,len(doc.tables)):                                           #srds,第0个表格是哪个
        table = tables[i]
        if((table.cell(0,0).text != "序号" and table.cell(0,0).text != "序\n号") or (table.cell(1,1).text[0] == "实" and table.cell(1,1).text[1] == "验") ):
            continue
        else:
            for j in range(1,len(table.rows)):                                       
                Order = table.cell(j,0).text
                #split返回的是列表
                #Content_Origin = encode('gb2312')
                #Content_Origin = Content_Origin.decode('utf-8')                          
                Content = table.cell(j,1).text.split("\n")               
                if(Order == '1'):
                    count += 1
                if(count == len(Keys)):
                    return DicOrigin                
                DicOrigin[Keys[count]].update({Order:Content})                                                                      
    #print(DicOrigin)                              
    return DicOrigin
# Count 增加条件
# Out of range?

AllContent = SearchCourseContent(doc,AllCourse)

#TestDic = {"信息":{'1': ['第1章  计算机系统', '计算的功能与发展；', '计算机软件系统和硬件系统的构成；', '冯诺依曼模型，计算机工作原理，计算机的逻辑组成及各部分的功能，计算机硬件发展面临的问题和发展方向。 ', '计算机网络、数据与信息、计算思维简介；', '计算机系统的分层模型。']}}

Keys = list(AllContent.keys())
for i in range(0,len(Keys)):
    print(Keys[i])
    print(AllContent[Keys[i]])

#print (json.dumps(AllContent, ensure_ascii=False))