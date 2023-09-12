import docx
import re
from docx import Document

# 检查是否能够正常打开
path = "F:\\Desktop_From_C\\project\\教学大纲.docx"
# 相对路径 -> path = "..\教学大纲.docx"

''' 
副本？
查找文件中的点出现的位置
idx = path.rfind(".")
print(file1_name[0:idx])
print(file1_name[idx:])
利用切片添加副本文字
file = path[0:idx]+ " - 副本" + path[idx:]
'''

try:    
    doc = Document(path)
except Exception as e:
    print("Error!",e)
else:
    print("Success!")
    #计算检索表格数量   
    tables = doc.tables
    TableLen = len(tables)

#提取课程名称,加入字典并返回字典
def SearchCourseName(doc):
    Dic = {}
    Dic["Course"] = []
    num = 0
    pattern = re.compile(r"《(.+?)》教学大纲")
    for p in doc.paragraphs:
        str = p.text    
        list = re.findall(pattern,str)
        if list:
            Dic["Course"].append(list)
            num = num+1
            #print(list)
    return Dic

AllCourse = SearchCourseName(doc)
#print(AllCourse)

#提取教学内容，存入并返回字典
#先把一个课的所有内容存在一个字典，再把字典存在之前的大字典
def SearchCourseContent(doc,dic1):
    #pattern = re.compile(r"[1-9][0-9]*")
    DicCon = {}
     
    for i in range(0,50):
        table = tables[i]
        Order = 1
        for j in range(1,len(table.rows)):
            if(table.cell(0,0).text!="序号" and table.cell(0,0).text!="序\n号"):
                break;
            Txt = table.cell(j,1).text
            Content = Txt.split("\n")
            DicCon = {Order:Content}
            print(DicCon)
            Order +=1
    return DicCon

AllContent = SearchCourseContent(doc,AllCourse)
#print(AllContent)

'''
#合并--迭代
#碰到序号1就跳下一个（注意末尾小心越界）
def Combine(Dic1,Dic2):
    Dic = {}
    AllKey1 = Dic1.keys()                 # 1-->Course  2-->Content
    AllKey2 = Dic2.keys()
    flag = 0
    for i in range(0,len(AllKey1)-1):
        for j in range(flag,len(AllKey2)-1):
            Dic[AllKey1[i]] = Dic2[flag]
            if(flag+1!=len(AllKey2)-1 and AllKey2[flag+1]=="1"):                        #判断下一个是否为末尾/值是否为1
                flag += 1
                break;               
    return Dic


#Course_Content = Combine(AllCourse, AllContent)
#print(Course_Content)
'''

#知识点的唯一相同点--序号    
'''  
现在输入问题解决了
如何输出是一个问题
对于一个表格里有多个值也是一个问题
怎样专门做一个函数也是一个问题
函数传参可以是文件名
也可以是什么都没有
'''
